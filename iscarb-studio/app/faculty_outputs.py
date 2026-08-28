from __future__ import annotations

from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor as DocxRGB

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .models import Blueprint

# ISCARB Original Identity — Saudi Academic Engineering
GREEN = "#0C533D"
GREEN2 = "#1D8B56"
TEAL = "#0A353E"
PURPLE = "#563C7D"
GOLD = "#C4A24F"
INK = "#1D2921"
MUTED = "#657169"
SAND = "#F6F4EF"
LINE = "#DDE4DF"
WHITE = "#FFFFFF"
SOFT_GREEN = "#EEF8F1"
SOFT_PURPLE = "#F2EDF7"
SOFT_GOLD = "#FBF6E8"
SOFT_TEAL = "#EEF4F4"

PHASE = {
    "IFHAM": (PURPLE, SOFT_PURPLE),
    "MARIS": (GREEN2, SOFT_GREEN),
    "ATQAN": (GOLD, SOFT_GOLD),
    "MAYYIZ": (TEAL, SOFT_TEAL),
}


def _short(text: str, limit: int = 150) -> str:
    text = " ".join(str(text or "").split())
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"


def _readiness_for_unit(bp: Blueprint, unit_no: int) -> list[str]:
    refs: list[str] = []
    for r in bp.readiness_alignment:
        if unit_no in r.evidence_units or unit_no in {3, 16, 19, 20}:
            label = f"{r.sku}: {', '.join(r.slo_refs)} → {', '.join(r.klo_refs)}"
            if label not in refs:
                refs.append(label)
    return refs


# -----------------------------------------------------------------------------
# Detailed Deck PDF — designed for reading, not presenting
# -----------------------------------------------------------------------------

def _pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ISCARBTitle", parent=styles["Title"], fontName="Helvetica-Bold",
            fontSize=27, leading=30, textColor=colors.HexColor(INK), spaceAfter=8,
        ),
        "thesis": ParagraphStyle(
            "ISCARBThesis", parent=styles["BodyText"], fontName="Helvetica",
            fontSize=11.5, leading=17, textColor=colors.HexColor(MUTED), spaceAfter=8,
        ),
        "h1": ParagraphStyle(
            "ISCARBH1", parent=styles["Heading1"], fontName="Helvetica-Bold",
            fontSize=20, leading=23, textColor=colors.HexColor(INK), spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "ISCARBH2", parent=styles["Heading2"], fontName="Helvetica-Bold",
            fontSize=10.3, leading=13, textColor=colors.HexColor(INK), spaceBefore=3, spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "ISCARBBody", parent=styles["BodyText"], fontName="Helvetica",
            fontSize=9.1, leading=13.2, textColor=colors.HexColor(INK), spaceAfter=3,
        ),
        "small": ParagraphStyle(
            "ISCARBSmall", parent=styles["BodyText"], fontName="Helvetica",
            fontSize=7.4, leading=10, textColor=colors.HexColor(MUTED), spaceAfter=2,
        ),
        "question": ParagraphStyle(
            "ISCARBQuestion", parent=styles["BodyText"], fontName="Helvetica-Bold",
            fontSize=12.2, leading=16, textColor=colors.HexColor(PURPLE), spaceAfter=0,
        ),
        "center": ParagraphStyle(
            "ISCARBCenter", parent=styles["BodyText"], alignment=TA_CENTER,
            fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=colors.HexColor(INK),
        ),
    }


def _pdf_bullet(text: str, style) -> Paragraph:
    return Paragraph("• " + escape(str(text)), style)


def _pdf_channel(title: str, items: list[str], style, tint: str, accent: str):
    if not items:
        return None
    rows = [[Paragraph(f"<b>{escape(title)}</b>", style)]]
    for item in items:
        rows.append([_pdf_bullet(item, style)])
    table = Table(rows, colWidths=[178 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(tint)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(accent)),
        ("BOX", (0, 0), (-1, -1), 0.65, colors.HexColor(LINE)),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor(LINE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def _pdf_header_footer(canvas, doc):
    canvas.saveState()
    w, h = A4
    canvas.setStrokeColor(colors.HexColor(GOLD))
    canvas.setLineWidth(0.8)
    canvas.line(18 * mm, h - 14 * mm, w - 18 * mm, h - 14 * mm)
    canvas.setFont("Helvetica-Bold", 7.4)
    canvas.setFillColor(colors.HexColor(GREEN))
    canvas.drawString(18 * mm, h - 11 * mm, "ISCARB · DETAILED DECK")
    canvas.setFont("Helvetica", 7.2)
    canvas.setFillColor(colors.HexColor(MUTED))
    canvas.drawRightString(w - 18 * mm, 10 * mm, f"Page {doc.page}")
    canvas.setStrokeColor(colors.HexColor(LINE))
    canvas.line(18 * mm, 14 * mm, w - 18 * mm, 14 * mm)
    canvas.restoreState()


def export_detailed_pdf(bp: Blueprint, out: Path) -> Path:
    out = Path(out)
    s = _pdf_styles()
    doc = SimpleDocTemplate(
        str(out), pagesize=A4,
        rightMargin=16 * mm, leftMargin=16 * mm, topMargin=21 * mm, bottomMargin=19 * mm,
        title=bp.lecture_title, author="ISCARB Faculty Studio",
    )
    story = []

    # Cover / executive map
    story += [
        Spacer(1, 9 * mm),
        Paragraph("ISCARB DETAILED DECK", ParagraphStyle(
            "CoverKicker", parent=s["small"], textColor=colors.HexColor(PURPLE),
            fontName="Helvetica-Bold", fontSize=8.5, leading=10, spaceAfter=8,
        )),
        Paragraph(escape(bp.lecture_title), s["title"]),
        Paragraph(escape(bp.engineering_thesis), s["thesis"]),
    ]
    facts = [
        ["90 MIN", "20 UNITS", "SOURCE LOCKED", "EVIDENCE → DECISION"],
        ["One live lecture", "IFHAM → MAYYIZ", "P1 is authoritative", "Assurance before release"],
    ]
    ft = Table([[Paragraph(f"<b>{x}</b>", s["center"]) for x in facts[0]],
                [Paragraph(x, s["small"]) for x in facts[1]]], colWidths=[44.5 * mm] * 4)
    ft.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(SOFT_GREEN)),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(LINE)),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor(LINE)),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story += [Spacer(1, 5 * mm), ft, Spacer(1, 6 * mm)]

    story.append(Paragraph("Central engineering crisis", s["h2"]))
    story.append(Paragraph(escape(bp.central_engineering_crisis), s["body"]))
    story.append(Paragraph("Professional / ethical purpose", s["h2"]))
    story.append(Paragraph(escape(bp.named_ethical_purpose), s["body"]))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("Primary source bundle", s["h2"]))
    for source in bp.source_manifest:
        story.append(_pdf_bullet(source, s["small"]))
    story.append(PageBreak())

    # CLO map
    story += [Paragraph("Five measurable learning outcomes", s["h1"]),
              Paragraph("Each outcome is paired with observable evidence rather than treated as a decorative statement.", s["thesis"])]
    clo_rows = [[Paragraph("CLO", s["center"]), Paragraph("Capability", s["center"]), Paragraph("Evidence expected", s["center"])]]
    for c in bp.clOs:
        clo_rows.append([
            Paragraph(f"<b>{escape(c.id)}</b>", s["body"]),
            Paragraph(escape(c.statement), s["body"]),
            Paragraph(escape(c.evidence_expected), s["body"]),
        ])
    clo_table = Table(clo_rows, colWidths=[17 * mm, 82 * mm, 79 * mm], repeatRows=1)
    clo_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(TEAL)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(LINE)),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(LINE)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story += [clo_table, PageBreak()]

    # One designed page per Unit
    for u in bp.units:
        accent, tint = PHASE[u.phase]
        head = Table([
            [Paragraph(f"<b>UNIT {u.number:02d}</b>", s["center"]), Paragraph(f"<b>{escape(u.phase)}</b>", s["center"]), Paragraph(f"<b>{u.planned_minutes} MIN</b>", s["center"])],
        ], colWidths=[34 * mm, 110 * mm, 34 * mm])
        head.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(accent)),
            ("TEXTCOLOR", (0, 0), (-1, -1), colors.white if u.phase != "ATQAN" else colors.HexColor(INK)),
            ("BOX", (0, 0), (-1, -1), 0, colors.HexColor(accent)),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        story += [head, Spacer(1, 3 * mm), Paragraph(escape(u.title), s["h1"])]

        qbox = Table([[Paragraph("ENGINEERING QUESTION", s["small"]), Paragraph(escape(u.engineering_question), s["question"]) ]], colWidths=[38 * mm, 140 * mm])
        qbox.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(tint)),
            ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(accent)),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        story += [qbox, Spacer(1, 3 * mm)]

        for channel_title, items, channel_tint, channel_accent in [
            ("SOURCE-LOCKED TECHNICAL CONTENT", u.core_content, "#F7F8F7", GREEN),
            ("ISCARB DECISION / PEDAGOGY", u.pedagogy_content, SOFT_PURPLE, PURPLE),
            ("CONTEXTUAL ENRICHMENT", u.enrichment_content, SOFT_GOLD, GOLD),
        ]:
            ch = _pdf_channel(channel_title, list(items), s["body"], channel_tint, channel_accent)
            if ch is not None:
                story += [ch, Spacer(1, 2.5 * mm)]

        if u.scenario_assumptions:
            assumptions = Table([[Paragraph("ASSUMPTIONS", s["small"]), Paragraph(escape(" · ".join(u.scenario_assumptions)), s["body"]) ]], colWidths=[35 * mm, 143 * mm])
            assumptions.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(SAND)),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story += [assumptions, Spacer(1, 2.5 * mm)]

        action_rows = [
            [Paragraph("YOU TRY", s["small"]), Paragraph(escape(u.student_action), s["body"])],
            [Paragraph("EVIDENCE", s["small"]), Paragraph(escape(u.evidence or "Observable learner artifact / decision trace"), s["body"])],
            [Paragraph("TAKEAWAY", s["small"]), Paragraph(escape(u.takeaway), s["body"])],
            [Paragraph("SOURCE", s["small"]), Paragraph(escape(u.source_anchor or "ISCARB pedagogy — no technical source claim"), s["small"])],
        ]
        refs = _readiness_for_unit(bp, u.number)
        if refs:
            action_rows.append([Paragraph("READINESS", s["small"]), Paragraph(escape(" | ".join(refs)), s["small"])])
        action = Table(action_rows, colWidths=[35 * mm, 143 * mm])
        action.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(tint)),
            ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(accent)),
            ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor(LINE)),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(action)

        if u.number == 19:
            story += [Spacer(1, 3 * mm), Paragraph("Four-level capability rubric", s["h2"])]
            rub = [[Paragraph("Criterion", s["center"]), Paragraph("4", s["center"]), Paragraph("3", s["center"]), Paragraph("2", s["center"]), Paragraph("1", s["center"])]]
            for r in bp.rubric_criteria:
                rub.append([
                    Paragraph(escape(r.criterion), s["small"]),
                    Paragraph(escape(_short(r.distinguished, 80)), s["small"]),
                    Paragraph(escape(_short(r.ready, 80)), s["small"]),
                    Paragraph(escape(_short(r.developing, 80)), s["small"]),
                    Paragraph(escape(_short(r.not_yet_ready, 80)), s["small"]),
                ])
            rt = Table(rub, colWidths=[42 * mm, 34 * mm, 34 * mm, 34 * mm, 34 * mm], repeatRows=1)
            rt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(TEAL)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor(LINE)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(rt)

        story.append(PageBreak())

    doc.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
    return out


# -----------------------------------------------------------------------------
# Instructor Guide DOCX — 90-minute run of show
# -----------------------------------------------------------------------------

def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def _set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(28)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(19)
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 2"].font.size = Pt(12)


def _doc_badge(paragraph, text: str, color_hex: str):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = DocxRGB.from_string(color_hex.replace("#", ""))


def export_instructor_guide(bp: Blueprint, out: Path) -> Path:
    out = Path(out)
    doc = Document()
    _set_doc_defaults(doc)

    title = doc.add_heading(bp.lecture_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    _doc_badge(p, "ISCARB INSTRUCTOR GUIDE", PURPLE)
    p.add_run("   ·   90 minutes   ·   20 units   ·   source-locked").italic = True
    doc.add_paragraph(bp.engineering_thesis)
    doc.add_heading("Teaching intent", level=1)
    doc.add_paragraph("Move students from framing an incomplete engineering problem to making and defending a bounded professional decision. The Presenter Deck stays visually sparse; this guide carries the facilitation detail.")

    doc.add_heading("Before class — 5-minute setup", level=1)
    for item in [
        "Open the Presenter Preview and verify all 20 Units render correctly.",
        "Keep the primary source available for source checks during discussion.",
        "Prepare one shared place for student evidence: board, LMS, document, or team canvas.",
        "Do not reveal Unit 5's named principle before students make the prediction.",
        "Treat BLOCKED output as faculty-review material; only RELEASE may carry ISCARB Verified.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("90-minute run of show", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Unit", "Phase", "Min", "Teacher move", "Student does", "Evidence"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, TEAL)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = DocxRGB(255, 255, 255)
            run.bold = True
    for u in bp.units:
        cells = table.add_row().cells
        values = [
            f"{u.number:02d}", u.phase, str(u.planned_minutes),
            _short(u.engineering_question, 115), _short(u.student_action, 115), _short(u.evidence or u.takeaway, 100),
        ]
        for idx, val in enumerate(values):
            cells[idx].text = val
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        accent, tint = PHASE[u.phase]
        _set_cell_shading(cells[1], tint)
    doc.add_page_break()

    for u in bp.units:
        accent, tint = PHASE[u.phase]
        h = doc.add_heading(f"Unit {u.number:02d} · {u.phase} · {u.title}", level=1)
        if h.runs:
            h.runs[0].font.color.rgb = DocxRGB.from_string(accent.replace("#", ""))
        p = doc.add_paragraph()
        _doc_badge(p, f"{u.planned_minutes} MIN", accent)
        p.add_run("   ")
        _doc_badge(p, f"CLO {' · '.join(u.clo_ids)}", PURPLE)

        doc.add_heading("Ask first", level=2)
        doc.add_paragraph(u.engineering_question)
        if u.number == 5:
            doc.add_paragraph("Facilitation rule: collect predictions before revealing or naming the source principle/model.", style="Intense Quote")

        if u.core_content:
            doc.add_heading("Reveal / anchor in the source", level=2)
            for item in u.core_content:
                doc.add_paragraph(item, style="List Bullet")

        if u.pedagogy_content:
            doc.add_heading("Facilitation moves", level=2)
            for item in u.pedagogy_content:
                doc.add_paragraph(item, style="List Bullet")

        if u.enrichment_content:
            doc.add_heading("Contextual enrichment — keep provenance explicit", level=2)
            for item in u.enrichment_content:
                doc.add_paragraph(item, style="List Bullet")
            if u.enrichment_basis:
                doc.add_paragraph("Basis: " + " | ".join(u.enrichment_basis))

        doc.add_heading("Student action", level=2)
        doc.add_paragraph(u.student_action)
        doc.add_heading("Look / listen for", level=2)
        doc.add_paragraph(u.evidence or "An inspectable learner decision, artifact, or evidence trace.")
        doc.add_heading("Land the unit", level=2)
        doc.add_paragraph(u.takeaway)

        refs = _readiness_for_unit(bp, u.number)
        meta = doc.add_table(rows=1, cols=2)
        meta.style = "Table Grid"
        meta.cell(0, 0).text = "Source"
        meta.cell(0, 1).text = u.source_anchor or "ISCARB pedagogy — no technical source claim"
        _set_cell_shading(meta.cell(0, 0), tint)
        if refs:
            row = meta.add_row().cells
            row[0].text = "Readiness"
            row[1].text = " | ".join(refs)
            _set_cell_shading(row[0], tint)
        doc.add_page_break()

    doc.add_heading("Assessment rubric", level=1)
    rt = doc.add_table(rows=1, cols=5)
    rt.style = "Table Grid"
    for i, h in enumerate(["Criterion", "4 · Distinguished", "3 · Ready", "2 · Developing", "1 · Not yet"]):
        rt.cell(0, i).text = h
        _set_cell_shading(rt.cell(0, i), TEAL)
        for run in rt.cell(0, i).paragraphs[0].runs:
            run.font.color.rgb = DocxRGB(255, 255, 255)
            run.bold = True
    for r in bp.rubric_criteria:
        cells = rt.add_row().cells
        for i, val in enumerate([r.criterion, r.distinguished, r.ready, r.developing, r.not_yet_ready]):
            cells[i].text = val
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    doc.save(out)
    return out


# -----------------------------------------------------------------------------
# Student Activity Pack DOCX — no instructor answers
# -----------------------------------------------------------------------------

def export_student_pack(bp: Blueprint, out: Path) -> Path:
    out = Path(out)
    doc = Document()
    _set_doc_defaults(doc)
    doc.add_heading(bp.lecture_title, 0)
    p = doc.add_paragraph()
    _doc_badge(p, "ISCARB STUDENT ACTIVITY PACK", GREEN)
    p.add_run("   ·   think → decide → prove")
    doc.add_paragraph("Use this pack during the 90-minute lecture. Write decisions, assumptions, trade-offs, evidence, and uncertainties. It intentionally omits instructor explanations and model answers.")

    doc.add_heading("Your five targets", level=1)
    for c in bp.clOs:
        doc.add_paragraph(f"{c.id}: {c.statement}", style="List Bullet")

    for phase in ["IFHAM", "MARIS", "ATQAN", "MAYYIZ"]:
        accent, _ = PHASE[phase]
        h = doc.add_heading(phase, level=1)
        if h.runs:
            h.runs[0].font.color.rgb = DocxRGB.from_string(accent.replace("#", ""))
        for u in [x for x in bp.units if x.phase == phase]:
            doc.add_heading(f"Unit {u.number:02d} · {u.title} · {u.planned_minutes} min", level=2)
            doc.add_paragraph(u.engineering_question)
            doc.add_paragraph("YOUR ACTION", style="Caption")
            doc.add_paragraph(u.student_action)
            if u.scenario_assumptions:
                doc.add_paragraph("Given assumptions: " + " | ".join(u.scenario_assumptions))
            prompt = doc.add_paragraph()
            prompt.add_run("Decision / response: ").bold = True
            prompt.add_run("________________________________________________________________________________")
            prompt = doc.add_paragraph()
            prompt.add_run("Evidence / reasoning: ").bold = True
            prompt.add_run("_______________________________________________________________________________")
            if u.number in {8, 9, 10, 17, 18, 20}:
                prompt = doc.add_paragraph()
                prompt.add_run("What would change your decision? ").bold = True
                prompt.add_run("________________________________________________________________________")

    doc.add_page_break()
    doc.add_heading("Portfolio evidence checklist", level=1)
    unit16 = bp.units[15]
    doc.add_paragraph(unit16.student_action)
    for item in [
        "Problem framing and explicit assumptions",
        "First-principles mechanism reasoning",
        "At least two defensible alternatives",
        "Explicit trade-off judgment",
        "Risk / uncertainty and what is monitored",
        "Evidence that could falsify or revise the decision",
        "Constraint mutation and adaptation",
        "Professional accountability / ethical purpose",
        "Readiness trace where supported",
        "Bounded final decision: approve / conditional / redesign / reject",
    ]:
        doc.add_paragraph("☐ " + item)

    doc.add_heading("Rubric — what strong work looks like", level=1)
    rt = doc.add_table(rows=1, cols=5)
    rt.style = "Table Grid"
    for i, h in enumerate(["Criterion", "4", "3", "2", "1"]):
        rt.cell(0, i).text = h
        _set_cell_shading(rt.cell(0, i), TEAL)
        for run in rt.cell(0, i).paragraphs[0].runs:
            run.font.color.rgb = DocxRGB(255, 255, 255)
            run.bold = True
    for r in bp.rubric_criteria:
        cells = rt.add_row().cells
        for i, val in enumerate([r.criterion, r.distinguished, r.ready, r.developing, r.not_yet_ready]):
            cells[i].text = _short(val, 145)
    doc.save(out)
    return out
